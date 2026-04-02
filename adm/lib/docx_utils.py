"""python-docx 서식 헬퍼.

보고서, 분석 문서 등을 python-docx로 생성할 때 공통으로 사용하는
테이블, 코드 블록, 불릿, 강조 등의 서식 유틸리티.

사용법:
    from docx import Document
    from adm.lib.docx_utils import add_styled_table, add_code_block, add_critical

    doc = Document()
    add_styled_table(doc, ["이름", "값"], [["a", "1"], ["b", "2"]])
    add_code_block(doc, "print('hello')")
    add_critical(doc, "경고", "이것은 중요한 문제입니다.")
    doc.save("report.docx")
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_cell_shading(cell, color_hex: str) -> None:
    """테이블 셀 배경색 설정.

    Args:
        cell: python-docx 테이블 셀 객체.
        color_hex: 6자리 hex 색상 코드 (예: "1B3A5C").
    """
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    shading.append(shd)


def add_styled_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
    header_bg: str = "1B3A5C",
    stripe_bg: str = "F2F6FC",
) -> None:
    """헤더 행 + 줄무늬 데이터 행으로 구성된 테이블 추가.

    Args:
        doc: python-docx Document 객체.
        headers: 헤더 텍스트 목록.
        rows: 데이터 행 목록 (각 행은 문자열 리스트).
        col_widths: 열 너비(cm) 목록. None이면 자동.
        header_bg: 헤더 배경 hex 색상 (기본: 남색).
        stripe_bg: 짝수행 배경 hex 색상 (기본: 연파랑).
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_bg)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, stripe_bg)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    doc.add_paragraph("")


def add_code_block(doc: Document, text: str) -> None:
    """회색 배경의 고정폭 폰트 코드 블록 추가.

    Args:
        doc: python-docx Document 객체.
        text: 코드 텍스트.
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "F5F5F5"},
    )
    pPr.append(shd)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    """불릿 포인트 추가. bold_prefix 부분은 볼드 처리.

    Args:
        doc: python-docx Document 객체.
        text: 본문 텍스트.
        bold_prefix: 볼드로 표시할 앞부분 텍스트.
    """
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)


def add_numbered(doc: Document, text: str) -> None:
    """번호 목록 항목 추가.

    Args:
        doc: python-docx Document 객체.
        text: 목록 항목 텍스트.
    """
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10)


def add_file_ref(doc: Document, path: str, lines: str = "") -> None:
    """파일 경로 참조 표시 (이탤릭, 회색).

    Args:
        doc: python-docx Document 객체.
        path: 파일 경로.
        lines: 라인 번호 범위 (예: "33-148").
    """
    p = doc.add_paragraph()
    run = p.add_run(f"📄 {path}" + (f" (라인 {lines})" if lines else ""))
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_critical(doc: Document, label: str, text: str) -> None:
    """빨간색 경고 강조 단락 추가.

    Args:
        doc: python-docx Document 객체.
        label: 강조 레이블 (예: "핵심 문제").
        text: 설명 텍스트.
    """
    p = doc.add_paragraph()
    run = p.add_run(f"⚠ {label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
